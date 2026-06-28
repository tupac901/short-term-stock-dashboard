from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "superpowers" / "specs" / "2026-06-28-stock-analysis-dashboard-design.md"
OUTPUT = ROOT / "docs" / "superpowers" / "specs" / "2026-06-28-stock-analysis-dashboard-design.docx"


def add_paragraph(document, line):
    if line.startswith("- "):
        document.add_paragraph(line[2:], style="List Bullet")
        return

    if line and line[0].isdigit() and ". " in line[:4]:
        number, text = line.split(". ", 1)
        if number.isdigit():
            document.add_paragraph(text, style="List Number")
            return

    document.add_paragraph(line)


def main():
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(13)

    title_added = False
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            paragraph = document.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_added = True
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif title_added:
            add_paragraph(document, line)
        else:
            document.add_paragraph(line)

    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
